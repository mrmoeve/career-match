from io import BytesIO
from xml.sax.saxutils import escape

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def extract_resume_candidate_name(resume_text: str) -> str:
    for line in (resume_text or "").splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if "@" in stripped or "|" in stripped:
            continue
        return stripped
    return "Candidate"


def build_optimized_resume_docx(optimized_resume_text: str) -> bytes:
    document = Document()
    for line in (optimized_resume_text or "").splitlines():
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
        elif stripped.isupper() and len(stripped.split()) <= 4:
            document.add_heading(stripped.title(), level=1)
        else:
            document.add_paragraph(stripped)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_optimized_resume_pdf(optimized_resume_text: str) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    for line in (optimized_resume_text or "").splitlines():
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 8))
        elif stripped.isupper() and len(stripped.split()) <= 4:
            story.append(Paragraph(escape(stripped.title()), styles["Heading1"]))
            story.append(Spacer(1, 8))
        else:
            story.append(Paragraph(escape(stripped), styles["BodyText"]))
            story.append(Spacer(1, 6))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def build_full_report_docx(package: dict) -> bytes:
    analysis = package["analysis"]
    generated = package["generated"]
    interview_dashboard = generated.get("interview_dashboard", {})
    career_coach = generated.get("career_coach", {})

    document = Document()
    document.add_heading("Career Match Export", level=0)
    document.add_paragraph(f"Created: {package['created_at']}")
    document.add_paragraph(f"Target role: {analysis['job_title']}")
    document.add_paragraph(f"Target company: {analysis['company_name']}")
    if package.get("job_url"):
        document.add_paragraph(f"Job URL: {package['job_url']}")
    if package.get("job_location"):
        document.add_paragraph(f"Location: {package['job_location']}")
    if package.get("job_date_posted"):
        document.add_paragraph(f"Date Posted: {package['job_date_posted']}")
    if package.get("job_category"):
        document.add_paragraph(f"Category: {package['job_category']}")
    document.add_paragraph(f"ATS score: {analysis['ats_score']}/100")
    if analysis.get("job_fit"):
        document.add_paragraph(f"Job fit score: {analysis['job_fit']['fit_score']}/100")
        document.add_paragraph(f"Recommendation: {analysis['job_fit']['recommendation']}")
        document.add_paragraph("Job fit reasoning:")
        for item in analysis["job_fit"].get("reasoning", []):
            document.add_paragraph(item, style="List Bullet")

    document.add_heading("Professional Summary", level=1)
    document.add_paragraph(generated["professional_summary"])

    document.add_heading("Tailored Resume Bullets", level=1)
    for bullet in generated["tailored_resume_bullets"]:
        document.add_paragraph(bullet, style="List Bullet")

    document.add_heading("Cover Letter", level=1)
    document.add_paragraph(generated["cover_letter"])

    document.add_heading("LinkedIn Recruiter Message", level=1)
    document.add_paragraph(generated["linkedin_recruiter_message"])

    if career_coach:
        document.add_heading("Career Coach", level=1)
        document.add_paragraph(career_coach.get("overview", ""))

        document.add_paragraph("Missing Skills")
        for item in career_coach.get("missing_skills", []):
            document.add_paragraph(item, style="List Bullet")

        document.add_paragraph("Missing Certifications")
        for item in career_coach.get("missing_certifications", []):
            document.add_paragraph(item, style="List Bullet")

        document.add_paragraph("Missing Technologies")
        for item in career_coach.get("missing_technologies", []):
            document.add_paragraph(item, style="List Bullet")

        document.add_paragraph("Missing Industry Experience")
        for item in career_coach.get("missing_industry_experience", []):
            document.add_paragraph(item, style="List Bullet")

        for heading, key, title_key in [
            ("30-Day Improvement Plan", "thirty_day_plan", "action"),
            ("90-Day Improvement Plan", "ninety_day_plan", "action"),
            ("Recommended Certifications", "recommended_certifications", "name"),
            ("Recommended Courses", "recommended_courses", "name"),
            ("Resume Improvements", "resume_improvements", "change"),
        ]:
            document.add_paragraph(heading)
            for item in career_coach.get(key, []):
                document.add_paragraph(item.get(title_key, ""), style="List Bullet")
                if item.get("why_it_matters"):
                    document.add_paragraph(f"Why: {item['why_it_matters']}")
                if item.get("reason"):
                    document.add_paragraph(f"Reason: {item['reason']}")
                document.add_paragraph(
                    f"Estimated Job Fit Score increase: +{item.get('estimated_job_fit_increase', 0)}"
                )

    document.add_heading("Interview Questions and Answers", level=1)
    for item in generated["interview_questions_and_answers"]:
        document.add_paragraph(f"Q: {item['question']}")
        document.add_paragraph(f"A: {item['answer']}")

    if interview_dashboard:
        document.add_heading("Interview Intelligence", level=1)
        document.add_paragraph(interview_dashboard.get("overview", ""))

        document.add_paragraph("Top 25 Likely Interview Questions")
        for item in interview_dashboard.get("top_25_likely_questions", []):
            document.add_paragraph(f"Q: {item['question']}")
            document.add_paragraph(f"Confidence: {item.get('confidence_score', 0)}/100")
            document.add_paragraph(f"Answer: {item.get('answer_summary', '')}")
            star = item.get("star_answer", {})
            document.add_paragraph(f"Situation: {star.get('situation', '')}")
            document.add_paragraph(f"Task: {star.get('task', '')}")
            document.add_paragraph(f"Action: {star.get('action', '')}")
            document.add_paragraph(f"Result: {star.get('result', '')}")

        document.add_paragraph("Technical Questions")
        for item in interview_dashboard.get("technical_questions", []):
            document.add_paragraph(f"Q: {item['question']}")
            document.add_paragraph(f"Confidence: {item.get('confidence_score', 0)}/100")
            document.add_paragraph(f"Answer: {item.get('answer_summary', '')}")

        document.add_paragraph("Behavioral Questions")
        for item in interview_dashboard.get("behavioral_questions", []):
            document.add_paragraph(f"Q: {item['question']}")
            document.add_paragraph(f"Confidence: {item.get('confidence_score', 0)}/100")
            document.add_paragraph(f"Answer: {item.get('answer_summary', '')}")

        document.add_paragraph("Questions to Ask the Interviewer")
        for item in interview_dashboard.get("questions_for_interviewer", []):
            document.add_paragraph(item, style="List Bullet")

        document.add_paragraph("Potential Challenges")
        for item in interview_dashboard.get("potential_challenges", []):
            document.add_paragraph(item.get("challenge", ""), style="List Bullet")
            document.add_paragraph(f"Why: {item.get('why_it_may_come_up', '')}")
            document.add_paragraph(f"Suggested response: {item.get('suggested_response', '')}")
            document.add_paragraph(f"Confidence: {item.get('confidence_score', 0)}/100")

    document.add_heading("Thank-You Email", level=1)
    document.add_paragraph(generated["thank_you_email"])
    if generated.get("results_summary_email"):
        document.add_heading("Results Summary Email", level=1)
        document.add_paragraph(generated["results_summary_email"])

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_full_report_pdf(package: dict) -> bytes:
    analysis = package["analysis"]
    generated = package["generated"]
    interview_dashboard = generated.get("interview_dashboard", {})
    career_coach = generated.get("career_coach", {})

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    def add_paragraph(text: str, style_name: str = "BodyText") -> None:
        safe_text = escape(text).replace("\n", "<br/>")
        story.append(Paragraph(safe_text, styles[style_name]))
        story.append(Spacer(1, 10))

    add_paragraph("Career Match Export", "Title")
    add_paragraph(f"Target role: {analysis['job_title']}")
    add_paragraph(f"Target company: {analysis['company_name']}")
    if package.get("job_url"):
        add_paragraph(f"Job URL: {package['job_url']}")
    if package.get("job_location"):
        add_paragraph(f"Location: {package['job_location']}")
    if package.get("job_date_posted"):
        add_paragraph(f"Date Posted: {package['job_date_posted']}")
    if package.get("job_category"):
        add_paragraph(f"Category: {package['job_category']}")
    add_paragraph(f"ATS score: {analysis['ats_score']}/100")
    if analysis.get("job_fit"):
        add_paragraph(f"Job fit score: {analysis['job_fit']['fit_score']}/100")
        add_paragraph(f"Recommendation: {analysis['job_fit']['recommendation']}")
        add_paragraph("Job fit reasoning", "Heading2")
        for item in analysis["job_fit"].get("reasoning", []):
            add_paragraph(f"- {item}")

    add_paragraph("Professional Summary", "Heading1")
    add_paragraph(generated["professional_summary"])

    add_paragraph("Tailored Resume Bullets", "Heading1")
    for bullet in generated["tailored_resume_bullets"]:
        add_paragraph(f"- {bullet}")

    add_paragraph("Cover Letter", "Heading1")
    add_paragraph(generated["cover_letter"])

    add_paragraph("LinkedIn Recruiter Message", "Heading1")
    add_paragraph(generated["linkedin_recruiter_message"])

    if career_coach:
        add_paragraph("Career Coach", "Heading1")
        add_paragraph(career_coach.get("overview", ""))
        for heading, key in [
            ("Missing Skills", "missing_skills"),
            ("Missing Certifications", "missing_certifications"),
            ("Missing Technologies", "missing_technologies"),
            ("Missing Industry Experience", "missing_industry_experience"),
        ]:
            add_paragraph(heading, "Heading2")
            for item in career_coach.get(key, []):
                add_paragraph(f"- {item}")
        for heading, key, title_key in [
            ("30-Day Improvement Plan", "thirty_day_plan", "action"),
            ("90-Day Improvement Plan", "ninety_day_plan", "action"),
            ("Recommended Certifications", "recommended_certifications", "name"),
            ("Recommended Courses", "recommended_courses", "name"),
            ("Resume Improvements", "resume_improvements", "change"),
        ]:
            add_paragraph(heading, "Heading2")
            for item in career_coach.get(key, []):
                add_paragraph(item.get(title_key, ""))
                if item.get("why_it_matters"):
                    add_paragraph(f"Why: {item['why_it_matters']}")
                if item.get("reason"):
                    add_paragraph(f"Reason: {item['reason']}")
                add_paragraph(f"Estimated Job Fit Score increase: +{item.get('estimated_job_fit_increase', 0)}")

    add_paragraph("Interview Questions and Answers", "Heading1")
    for item in generated["interview_questions_and_answers"]:
        add_paragraph(f"Q: {item['question']}")
        add_paragraph(f"A: {item['answer']}")

    if interview_dashboard:
        add_paragraph("Interview Intelligence", "Heading1")
        add_paragraph(interview_dashboard.get("overview", ""))

        add_paragraph("Top 25 Likely Interview Questions", "Heading2")
        for item in interview_dashboard.get("top_25_likely_questions", []):
            add_paragraph(f"Q: {item['question']}")
            add_paragraph(f"Confidence: {item.get('confidence_score', 0)}/100")
            add_paragraph(f"Answer: {item.get('answer_summary', '')}")
            star = item.get("star_answer", {})
            add_paragraph(f"Situation: {star.get('situation', '')}")
            add_paragraph(f"Task: {star.get('task', '')}")
            add_paragraph(f"Action: {star.get('action', '')}")
            add_paragraph(f"Result: {star.get('result', '')}")

        add_paragraph("Technical Questions", "Heading2")
        for item in interview_dashboard.get("technical_questions", []):
            add_paragraph(f"Q: {item['question']}")
            add_paragraph(f"Confidence: {item.get('confidence_score', 0)}/100")
            add_paragraph(f"Answer: {item.get('answer_summary', '')}")

        add_paragraph("Behavioral Questions", "Heading2")
        for item in interview_dashboard.get("behavioral_questions", []):
            add_paragraph(f"Q: {item['question']}")
            add_paragraph(f"Confidence: {item.get('confidence_score', 0)}/100")
            add_paragraph(f"Answer: {item.get('answer_summary', '')}")

        add_paragraph("Questions to Ask the Interviewer", "Heading2")
        for item in interview_dashboard.get("questions_for_interviewer", []):
            add_paragraph(f"- {item}")

        add_paragraph("Potential Challenges", "Heading2")
        for item in interview_dashboard.get("potential_challenges", []):
            add_paragraph(item.get("challenge", ""))
            add_paragraph(f"Why: {item.get('why_it_may_come_up', '')}")
            add_paragraph(f"Suggested response: {item.get('suggested_response', '')}")
            add_paragraph(f"Confidence: {item.get('confidence_score', 0)}/100")

    add_paragraph("Thank-You Email", "Heading1")
    add_paragraph(generated["thank_you_email"])
    if generated.get("results_summary_email"):
        add_paragraph("Results Summary Email", "Heading1")
        add_paragraph(generated["results_summary_email"])

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()
